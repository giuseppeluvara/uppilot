"""Generazione della bozza in formato Word (.docx) — §7.

Struttura (curata, professionale): intestazione, "In fatto", "In diritto"
(motivazione per ciascuna domanda con onere probatorio, non contestazioni,
allegati e quesiti), P.Q.M.

Lo stile tipografico è definito nel CODICE (font, gerarchia, giustificazione,
riquadri) così da essere versionabile e riproducibile. Resta il gancio per un
template.docx fornito dall'utente: se presente, si rispetta quello.

Principio §1: i punti discrezionali restano quesiti, riportati come callout
"Da decidere — verifica tu" visivamente distinti — mai conclusioni definitive.
"""
from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path

from django.conf import settings
from docx import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from apps.casi.privacy import maschera_residui

from .models import Bozza, EventoDecisionale, FattoProcessuale
from .serializers import FattoProcessualeSerializer

_PARTE_LABEL = {"attore": "attore", "convenuto": "convenuto/ricorrente"}

_FONT = "Times New Roman"
_ACCENTO = RGBColor(0x1F, 0x3A, 0x5F)  # blu ardesia sobrio (unico accento)
_QUESITO_FILL = "FFF3CD"  # ambra chiaro per i callout "Da decidere"
_QUESITO_BORDO = "E0A800"
_QUESITO_LABEL = RGBColor(0x8A, 0x6D, 0x3B)
_AVVISO_FILL = "F1F1F1"
_ROSSO = RGBColor(0xB0, 0x00, 0x20)


def _nome_file(percorso: str) -> str:
    return percorso.split("/")[-1]


def _documento_base() -> tuple["DocxDocument", bool]:
    """Usa template.docx dei sample se presente; restituisce (doc, da_template)."""
    template: Path = settings.SAMPLE_OUTPUT_DIR / "template.docx"
    if template.exists():
        return DocxDocument(str(template)), True
    return DocxDocument(), False


def _ombreggia(paragrafo, fill_hex: str) -> None:
    """Imposta lo sfondo (shading) di un paragrafo."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    paragrafo._p.get_or_add_pPr().append(shd)


def _bordo_sinistro(paragrafo, color_hex: str) -> None:
    """Aggiunge una barra di bordo a sinistra (effetto callout)."""
    pPr = paragrafo._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color_hex)
    pbdr.append(left)
    pPr.append(pbdr)


def _assicura_stile_quesito(doc) -> None:
    if "Quesito" not in [s.name for s in doc.styles]:
        st = doc.styles.add_style("Quesito", WD_STYLE_TYPE.PARAGRAPH)
        st.base_style = doc.styles["Normal"]
        st.font.size = Pt(11)
        st.font.italic = True
        st.paragraph_format.left_indent = Pt(12)
        st.paragraph_format.space_before = Pt(4)
        st.paragraph_format.space_after = Pt(4)


def _applica_stili(doc) -> None:
    """Stile professionale: corpo serif giustificato, gerarchia titoli con accento."""
    normal = doc.styles["Normal"]
    normal.font.name = _FONT
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_after = Pt(6)
    pf.line_spacing = 1.3

    for nome, dim in (("Title", 20), ("Heading 1", 14), ("Heading 2", 12.5)):
        st = doc.styles[nome]
        st.font.name = _FONT
        st.font.size = Pt(dim)
        st.font.bold = True
        st.font.color.rgb = _ACCENTO
    h1 = doc.styles["Heading 1"].paragraph_format
    h1.space_before = Pt(14)
    h1.space_after = Pt(4)


def _numeri_pagina(doc) -> None:
    """Footer centrato con numero di pagina."""
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Pag. ")
    run = p.add_run()
    inizio = OxmlElement("w:fldChar")
    inizio.set(qn("w:fldCharType"), "begin")
    istr = OxmlElement("w:instrText")
    istr.set(qn("xml:space"), "preserve")
    istr.text = "PAGE"
    fine = OxmlElement("w:fldChar")
    fine.set(qn("w:fldCharType"), "end")
    run._r.append(inizio)
    run._r.append(istr)
    run._r.append(fine)


def _depseudonimizza(testo: str, mappa: dict) -> str:
    """Ri-sostituisce i placeholder canonici con i valori reali (export in chiaro)."""
    if not testo or not mappa:
        return testo
    # Placeholder più lunghi prima, per evitare collisioni di prefisso.
    for ph in sorted(mappa, key=len, reverse=True):
        testo = testo.replace(ph, mappa[ph])
    return _pulizia_spazi_chiaro(testo)


def _pulizia_spazi_chiaro(testo: str) -> str:
    if not testo:
        return testo
    testo = re.sub(r"\b(via|viale|piazza|corso|largo)([A-ZÀ-Ö])", r"\1 \2", testo)
    testo = re.sub(r"\s+([,.;:!?])", r"\1", testo)
    testo = re.sub(r"[ \t]{2,}", " ", testo)
    return testo


def genera_docx(lavoro, in_chiaro: bool = False) -> bytes:
    doc, da_template = _documento_base()
    if not da_template:
        _applica_stili(doc)
    _assicura_stile_quesito(doc)

    mappa = lavoro.mappa_entita or {}

    def chiaro(t):
        if in_chiaro:
            return _depseudonimizza(t or "", mappa)
        return maschera_residui(t or "", mappa)

    # --- Intestazione ---
    titolo = doc.add_paragraph("BOZZA DI PROVVEDIMENTO", style="Title")
    titolo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = sub.add_run(chiaro(lavoro.titolo) if in_chiaro else f"Fascicolo #{lavoro.id}")
    run_sub.italic = True

    avviso = doc.add_paragraph()
    _ombreggia(avviso, _AVVISO_FILL)
    nota = avviso.add_run(
        "Bozza assistita, soggetta a revisione umana. I passaggi contrassegnati "
        "[DA DECIDERE] richiedono la valutazione dell'operatore."
    )
    nota.italic = True
    nota.font.size = Pt(10)
    if in_chiaro:
        warn = avviso.add_run(
            " ATTENZIONE: questo documento contiene i DATI PERSONALI REALI delle parti "
            "(versione in chiaro). Trattare con le dovute cautele."
        )
        warn.italic = True
        warn.bold = True
        warn.font.size = Pt(10)
        warn.font.color.rgb = _ROSSO

    # --- In fatto ---
    doc.add_heading("In fatto", level=1)
    bozza = Bozza.objects.filter(lavoro=lavoro).first()
    doc.add_paragraph(chiaro(bozza.in_fatto) if bozza and bozza.in_fatto else "—")

    # --- In diritto (motivazione per ciascuna domanda) ---
    doc.add_heading("In diritto", level=1)
    richieste = list(lavoro.richieste.all())
    if not richieste:
        doc.add_paragraph("Nessuna richiesta analizzata.")
    for i, r in enumerate(richieste, start=1):
        parte = _PARTE_LABEL.get(r.parte_richiedente, r.parte_richiedente)
        doc.add_heading(f"{i}. Domanda di parte {parte}", level=2)
        doc.add_paragraph(chiaro(r.testo))
        meta = doc.add_paragraph()
        meta.add_run("Tipo: ").bold = True
        meta.add_run(getattr(r, "get_tipo_display", lambda: r.tipo)())
        meta.add_run(" · Confidenza: ").bold = True
        meta.add_run(f"{round((r.confidence or 0) * 100)}%")

        for flag in r.flags or []:
            fp = doc.add_paragraph(style="Quesito")
            _ombreggia(fp, _QUESITO_FILL)
            _bordo_sinistro(fp, _QUESITO_BORDO)
            label = fp.add_run("Da rivedere: ")
            label.bold = True
            label.font.color.rgb = _QUESITO_LABEL
            fp.add_run(chiaro(str(flag)))

        if r.motivazione:
            doc.add_paragraph(chiaro(r.motivazione))

        if r.onere_probatorio:
            p = doc.add_paragraph()
            p.add_run("Onere probatorio: ").bold = True
            p.add_run(chiaro(r.onere_probatorio))

        if r.fonti_tracciate:
            p = doc.add_paragraph()
            p.add_run("Fonti interne tracciate:").bold = True
            for fonte in (r.fonti_tracciate or [])[:3]:
                if not isinstance(fonte, dict):
                    continue
                nome = fonte.get("documento_nome") if in_chiaro else f"Documento {fonte.get('documento_id')}"
                score = int(round(float(fonte.get("score") or 0) * 100))
                label = fonte.get("affidabilita_label") or fonte.get("affidabilita") or "Riscontro"
                testo_fonte = f"{nome} — {label} {score}%"
                snippet = fonte.get("snippet")
                if snippet:
                    testo_fonte += f": {snippet}"
                doc.add_paragraph(chiaro(testo_fonte), style="List Bullet")

        if r.non_contestazioni:
            p = doc.add_paragraph()
            p.add_run("Non contestazioni:").bold = True
            for nc in r.non_contestazioni:
                doc.add_paragraph(chiaro(str(nc)), style="List Bullet")

        allegati = list(r.allegati_collegati.all())
        if allegati:
            p = doc.add_paragraph()
            p.add_run("Allegati collegati: ").bold = True
            p.add_run(
                ", ".join(
                    _nome_file(a.file.name) if in_chiaro else f"Documento {a.id}"
                    for a in allegati
                )
            )

        for q in r.quesiti_aperti:
            qp = doc.add_paragraph(style="Quesito")
            _ombreggia(qp, _QUESITO_FILL)
            _bordo_sinistro(qp, _QUESITO_BORDO)
            label = qp.add_run("Da decidere — verifica tu: ")
            label.bold = True
            label.font.color.rgb = _QUESITO_LABEL
            qp.add_run(chiaro(q))

    # --- P.Q.M. ---
    doc.add_heading("P.Q.M.", level=1)
    if bozza and bozza.pqm:
        doc.add_paragraph(chiaro(bozza.pqm))
    else:
        doc.add_paragraph("[Da compilare dall'operatore]")

    if not da_template:
        _numeri_pagina(doc)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def genera_audit_docx(lavoro) -> bytes:
    """Allegato interno di audit: matrice, fonti, lacune e scelte umane."""

    doc, da_template = _documento_base()
    if not da_template:
        _applica_stili(doc)

    doc.add_paragraph("ALLEGATO DI AUDIT - UPPILOT", style="Title").alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )
    doc.add_paragraph(f"Fascicolo #{lavoro.id} - {lavoro.titolo}")

    avviso = doc.add_paragraph()
    _ombreggia(avviso, _AVVISO_FILL)
    run = avviso.add_run(
        "Documento interno di controllo: riepiloga fonti, lacune e scelte umane. "
        "Non sostituisce la revisione del magistrato/operatore."
    )
    run.italic = True
    run.font.size = Pt(10)

    doc.add_heading("1. Matrice richieste/prove", level=1)
    righe = (
        FattoProcessuale.objects.filter(richiesta__lavoro=lavoro)
        .select_related("richiesta", "richiesta__lavoro")
        .prefetch_related("richiesta__allegati_collegati")
        .order_by("richiesta__ordine", "ordine", "id")
    )
    if not righe.exists():
        doc.add_paragraph("Nessuna riga matrice disponibile.")
    for i, fatto in enumerate(righe, start=1):
        data = FattoProcessualeSerializer(fatto).data
        doc.add_heading(f"{i}. Richiesta {data['richiesta_id']}", level=2)
        doc.add_paragraph(data["richiesta_testo"] or "—")

        stato = doc.add_paragraph()
        stato.add_run("Stato prova: ").bold = True
        stato.add_run(data["stato_prova_label"])
        stato.add_run(" · Contraddittorio: ").bold = True
        stato.add_run(data["stato_contraddittorio_label"])
        stato.add_run(" · Suggerito: ").bold = True
        stato.add_run(data["stato_contraddittorio_suggerito_label"])

        if data["testo"]:
            p = doc.add_paragraph()
            p.add_run("Fatto rilevante: ").bold = True
            p.add_run(data["testo"])
        if data["note_operatore"]:
            p = doc.add_paragraph()
            p.add_run("Note operatore: ").bold = True
            p.add_run(data["note_operatore"])
        if data["note_contraddittorio"]:
            p = doc.add_paragraph()
            p.add_run("Note contraddittorio: ").bold = True
            p.add_run(data["note_contraddittorio"])

        if data["lacune"]:
            doc.add_paragraph("Lacune / punti da verificare:")
            for lacuna in data["lacune"]:
                doc.add_paragraph(str(lacuna), style="List Bullet")

        if data["fonti"]:
            doc.add_paragraph("Fonti principali:")
            for fonte in data["fonti"][:5]:
                score = int(round(float(fonte.get("score") or 0) * 100))
                testo = (
                    f"Documento {fonte.get('documento_id')} - "
                    f"{fonte.get('sezione_label', '')} - {score}%"
                )
                snippet = fonte.get("snippet")
                if snippet:
                    testo += f": {snippet}"
                doc.add_paragraph(testo, style="List Bullet")

    doc.add_heading("2. Registro decisionale", level=1)
    eventi = EventoDecisionale.objects.filter(lavoro=lavoro).select_related(
        "utente", "richiesta", "fatto"
    )[:40]
    if not eventi:
        doc.add_paragraph("Nessun evento decisionale registrato.")
    for evento in eventi:
        p = doc.add_paragraph()
        p.add_run(f"{evento.created_at:%d/%m/%Y %H:%M} - {evento.get_tipo_display()}: ").bold = True
        p.add_run(evento.descrizione or evento.campo or "evento")
        if evento.utente:
            p.add_run(f" ({evento.utente.username})")

    if not da_template:
        _numeri_pagina(doc)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
